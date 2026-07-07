# ui/reports.py
import customtkinter as ctk
from database.db import (
    get_daily_report, get_weekly_report, get_weekly_report_data,
    get_daily_balance, get_sales_by_payment_type, get_connection,
    get_business_date, get_weekly_start_date, get_recent_expenses,
    get_all_expenses, get_item_type
)
from datetime import datetime, timedelta
import csv
import os
from ui.font_utils import get_font, get_font_bold, get_font_size

# Try to import document generators
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("python-docx not installed. Run: pip install python-docx")

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    from reportlab.lib.units import inch
    from reportlab.pdfgen import canvas
    HAS_REPORTLAB = True
except ImportError:
    HAS_REPORTLAB = False
    print("reportlab not installed. Run: pip install reportlab")

class ReportsFrame(ctk.CTkFrame):
    def __init__(self, parent):
        super().__init__(parent)
        
        self.font_size = get_font_size()
        
        # Title
        title = ctk.CTkLabel(
            self,
            text="📈 Reports",
            font=get_font_bold(28)
        )
        title.pack(pady=(20, 10))
        
        # Date Selection Frame
        date_frame = ctk.CTkFrame(self)
        date_frame.pack(fill="x", padx=20, pady=5)
        
        ctk.CTkLabel(date_frame, text="Select Date:", font=get_font(self.font_size)).pack(side="left", padx=5)
        
        # Date entry with calendar button
        self.date_entry = ctk.CTkEntry(date_frame, placeholder_text="YYYY-MM-DD", width=150)
        self.date_entry.pack(side="left", padx=5)
        self.date_entry.insert(0, datetime.now().strftime("%Y-%m-%d"))
        
        ctk.CTkButton(
            date_frame,
            text="📅",
            width=35,
            command=self.open_calendar
        ).pack(side="left", padx=2)
        
        ctk.CTkButton(
            date_frame,
            text="📊 Generate Report",
            command=self.generate_report_for_date,
            width=150,
            fg_color="#1f538d"
        ).pack(side="left", padx=10)
        
        ctk.CTkButton(
            date_frame,
            text="📅 Today",
            command=self.go_to_today,
            width=80
        ).pack(side="left", padx=5)
        
        ctk.CTkButton(
            date_frame,
            text="◀ Yesterday",
            command=self.go_to_yesterday,
            width=100
        ).pack(side="left", padx=5)
        
        ctk.CTkButton(
            date_frame,
            text="Tomorrow ▶",
            command=self.go_to_tomorrow,
            width=100
        ).pack(side="left", padx=5)
        
        # Buttons
        btn_frame = ctk.CTkFrame(self)
        btn_frame.pack(fill="x", padx=20, pady=5)
        
        ctk.CTkButton(
            btn_frame,
            text="📊 Daily Report",
            command=self.show_daily_report,
            width=140
        ).pack(side="left", padx=3)
        
        ctk.CTkButton(
            btn_frame,
            text="📈 Weekly Report",
            command=self.show_weekly_report,
            width=140
        ).pack(side="left", padx=3)
        
        ctk.CTkButton(
            btn_frame,
            text="📊 Balance Report",
            command=self.show_balance_report,
            width=140
        ).pack(side="left", padx=3)
        
        ctk.CTkButton(
            btn_frame,
            text="📊 All Data",
            command=self.show_all_data_report,
            width=140
        ).pack(side="left", padx=3)
        
        # Export buttons
        export_frame = ctk.CTkFrame(btn_frame, fg_color="transparent")
        export_frame.pack(side="left", padx=5)
        
        ctk.CTkButton(
            export_frame,
            text="📥 CSV",
            command=self.export_csv,
            width=80,
            fg_color="green"
        ).pack(side="left", padx=2)
        
        if HAS_DOCX:
            ctk.CTkButton(
                export_frame,
                text="📄 Word",
                command=self.export_word,
                width=80,
                fg_color="#2b5797"
            ).pack(side="left", padx=2)
        
        if HAS_REPORTLAB:
            ctk.CTkButton(
                export_frame,
                text="📕 PDF",
                command=self.export_pdf,
                width=80,
                fg_color="#c0392b"
            ).pack(side="left", padx=2)
        
        # Report display
        self.report_frame = ctk.CTkScrollableFrame(self)
        self.report_frame.pack(fill="both", expand=True, padx=20, pady=10)
        
        # Show daily report by default
        self.show_daily_report()
    
    def open_calendar(self):
        """Open a calendar popup to select a date"""
        font_size = get_font_size()
        
        # Create popup window
        calendar_window = ctk.CTkToplevel(self)
        calendar_window.title("Select Date")
        calendar_window.geometry("300x280")
        calendar_window.transient(self)
        calendar_window.grab_set()
        calendar_window.focus_force()
        calendar_window.lift()
        
        # Center the window
        calendar_window.update_idletasks()
        width = 300
        height = 280
        x = (calendar_window.winfo_screenwidth() // 2) - (width // 2)
        y = (calendar_window.winfo_screenheight() // 2) - (height // 2)
        calendar_window.geometry(f'{width}x{height}+{x}+{y}')
        
        # Get current date from entry
        current_date_str = self.date_entry.get().strip()
        try:
            current_date = datetime.strptime(current_date_str, "%Y-%m-%d")
        except:
            current_date = datetime.now()
        
        # Try to use tkcalendar
        try:
            from tkcalendar import Calendar
            
            ctk.CTkLabel(
                calendar_window,
                text="Select a Date",
                font=get_font_bold(font_size + 2)
            ).pack(pady=10)
            
            # Create calendar widget
            cal = Calendar(
                calendar_window,
                selectmode='day',
                year=current_date.year,
                month=current_date.month,
                day=current_date.day,
                date_pattern='yyyy-mm-dd',
                background='#2b2b2b',
                foreground='white',
                selectbackground='#1f538d',
                selectforeground='white',
                weekendbackground='#3a3a3a',
                weekendforeground='gray',
                headersbackground='#1a1a1a',
                headersforeground='white'
            )
            cal.pack(pady=10, padx=10)
            
            def select_date():
                selected_date = cal.get_date()
                self.date_entry.delete(0, "end")
                self.date_entry.insert(0, selected_date)
                calendar_window.destroy()
                # Auto-generate report
                self.generate_report_for_date()
            
            btn_frame = ctk.CTkFrame(calendar_window)
            btn_frame.pack(pady=10)
            
            ctk.CTkButton(
                btn_frame,
                text="✅ Select",
                command=select_date,
                width=100
            ).pack(side="left", padx=5)
            
            ctk.CTkButton(
                btn_frame,
                text="Cancel",
                command=calendar_window.destroy,
                width=100
            ).pack(side="left", padx=5)
        except ImportError:
            # Fallback if tkcalendar not available
            ctk.CTkLabel(
                calendar_window,
                text="Please enter date manually:",
                font=get_font(font_size)
            ).pack(pady=10)
            
            entry = ctk.CTkEntry(calendar_window, placeholder_text="YYYY-MM-DD", width=200)
            entry.pack(pady=10)
            entry.insert(0, current_date.strftime("%Y-%m-%d"))
            
            def select_date():
                date_str = entry.get().strip()
                try:
                    datetime.strptime(date_str, "%Y-%m-%d")
                    self.date_entry.delete(0, "end")
                    self.date_entry.insert(0, date_str)
                    calendar_window.destroy()
                    self.generate_report_for_date()
                except ValueError:
                    error_label = ctk.CTkLabel(
                        calendar_window,
                        text="Invalid date format! Use YYYY-MM-DD",
                        text_color="red"
                    )
                    error_label.pack(pady=5)
            
            btn_frame = ctk.CTkFrame(calendar_window)
            btn_frame.pack(pady=10)
            
            ctk.CTkButton(
                btn_frame,
                text="✅ Select",
                command=select_date,
                width=100
            ).pack(side="left", padx=5)
            
            ctk.CTkButton(
                btn_frame,
                text="Cancel",
                command=calendar_window.destroy,
                width=100
            ).pack(side="left", padx=5)
    
    def go_to_today(self):
        """Go to today's date"""
        self.date_entry.delete(0, "end")
        self.date_entry.insert(0, datetime.now().strftime("%Y-%m-%d"))
        self.generate_report_for_date()
    
    def go_to_yesterday(self):
        """Go to yesterday's date"""
        yesterday = (datetime.now() - timedelta(days=1)).strftime("%Y-%m-%d")
        self.date_entry.delete(0, "end")
        self.date_entry.insert(0, yesterday)
        self.generate_report_for_date()
    
    def go_to_tomorrow(self):
        """Go to tomorrow's date"""
        tomorrow = (datetime.now() + timedelta(days=1)).strftime("%Y-%m-%d")
        self.date_entry.delete(0, "end")
        self.date_entry.insert(0, tomorrow)
        self.generate_report_for_date()
    
    def generate_report_for_date(self):
        """Generate report for the selected date"""
        date_str = self.date_entry.get().strip()
        try:
            # Validate date format
            selected_date = datetime.strptime(date_str, "%Y-%m-%d")
            
            # Check if date is in the future
            if selected_date > datetime.now():
                self.show_message("⚠️ Cannot generate report for future dates", "orange")
                return
            
            # Show daily report for selected date
            self.show_daily_report_for_date(date_str)
        except ValueError:
            self.show_message("❌ Invalid date format. Use YYYY-MM-DD", "red")
    
    def show_daily_report_for_date(self, date_str):
        """Show daily report for a specific date"""
        self.clear_report()
        
        font_size = get_font_size()
        report = get_daily_report(date_str)
        expenses = self.get_expenses_for_period(date_str, date_str)
        expense_summary = self.get_expenses_summary(date_str, date_str)
        sales_summary = self.get_sales_summary_with_types(date_str, date_str)
        
        # Title
        ctk.CTkLabel(
            self.report_frame,
            text="VALLEY VIEW MOTEL",
            font=get_font_bold(24),
            text_color="#1f538d"
        ).pack(pady=5)
        
        ctk.CTkLabel(
            self.report_frame,
            text=f"📊 Daily Report - {date_str}",
            font=get_font_bold(20)
        ).pack(pady=5)
        
        ctk.CTkLabel(
            self.report_frame,
            text=f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            font=get_font(font_size - 2),
            text_color="gray"
        ).pack(pady=2)
        
        ctk.CTkFrame(self.report_frame, height=2, fg_color="gray").pack(fill="x", pady=10)
        
        # Check if it's Saturday (Collection Day)
        if report.get('is_saturday', False):
            collection_frame = ctk.CTkFrame(self.report_frame, fg_color="#1a3a1a")
            collection_frame.pack(fill="x", pady=5)
            
            ctk.CTkLabel(
                collection_frame,
                text="💰 SATURDAY - CASH COLLECTION DAY 💰",
                font=get_font_bold(16),
                text_color="gold"
            ).pack(pady=5)
            
            ctk.CTkLabel(
                collection_frame,
                text=f"Cash Collected: UGX {report.get('cash_collected', 0):,.0f}",
                font=get_font(font_size)
            ).pack(pady=2)
            
            ctk.CTkLabel(
                collection_frame,
                text=f"Merchant Collected: UGX {report.get('merchant_collected', 0):,.0f}",
                font=get_font(font_size)
            ).pack(pady=2)
            
            ctk.CTkLabel(
                collection_frame,
                text=f"Total Collected: UGX {report.get('cash_collected', 0) + report.get('merchant_collected', 0):,.0f}",
                font=get_font_bold(font_size + 2),
                text_color="green"
            ).pack(pady=2)
        
        # Balance Summary
        balance_frame = ctk.CTkFrame(self.report_frame, fg_color="#1a2a3a")
        balance_frame.pack(fill="x", pady=5)
        
        ctk.CTkLabel(balance_frame, text="BALANCE SUMMARY", font=get_font_bold(font_size + 2)).pack(pady=5)
        
        balance_stats = [
            ("Opening Balance", f"UGX {report.get('opening_balance', 0):,.0f}"),
            ("Closing Balance", f"UGX {report.get('closing_balance', 0):,.0f}")
        ]
        
        for label, value in balance_stats:
            row = ctk.CTkFrame(balance_frame)
            row.pack(fill="x", pady=2)
            ctk.CTkLabel(row, text=label, font=get_font(font_size), width=200).pack(side="left", padx=20)
            ctk.CTkLabel(row, text=value, font=get_font_bold(font_size)).pack(side="right", padx=20)
        
        # Payment Type Breakdown
        payment_frame = ctk.CTkFrame(self.report_frame, fg_color="#1a2a3a")
        payment_frame.pack(fill="x", pady=5)
        
        ctk.CTkLabel(payment_frame, text="PAYMENT BREAKDOWN", font=get_font_bold(font_size + 2)).pack(pady=5)
        
        payment_stats = [
            ("Cash Sales", f"UGX {report.get('cash_sales', 0):,.0f}"),
            ("Merchant Sales", f"UGX {report.get('merchant_sales', 0):,.0f}")
        ]
        
        for label, value in payment_stats:
            row = ctk.CTkFrame(payment_frame)
            row.pack(fill="x", pady=2)
            ctk.CTkLabel(row, text=label, font=get_font(font_size), width=200).pack(side="left", padx=20)
            ctk.CTkLabel(row, text=value, font=get_font_bold(font_size)).pack(side="right", padx=20)
        
        # Sales Summary with Type Breakdown
        sales_frame = ctk.CTkFrame(self.report_frame, fg_color="#1a2a3a")
        sales_frame.pack(fill="x", pady=5)
        
        ctk.CTkLabel(sales_frame, text="SALES SUMMARY", font=get_font_bold(font_size + 2)).pack(pady=5)
        
        # Overall totals
        total_row = ctk.CTkFrame(sales_frame)
        total_row.pack(fill="x", pady=2)
        ctk.CTkLabel(total_row, text="Total Sales", font=get_font_bold(font_size), width=200).pack(side="left", padx=20)
        ctk.CTkLabel(total_row, text=f"UGX {sales_summary['total']['revenue']:,.0f} ({sales_summary['total']['count']} transactions)", font=get_font_bold(font_size)).pack(side="right", padx=20)
        
        # Inventory Sales
        inv_row = ctk.CTkFrame(sales_frame)
        inv_row.pack(fill="x", pady=2)
        ctk.CTkLabel(inv_row, text="📦 Inventory Sales", font=get_font(font_size), width=200).pack(side="left", padx=20)
        ctk.CTkLabel(inv_row, text=f"UGX {sales_summary['inventory']['revenue']:,.0f} ({sales_summary['inventory']['count']} items)", font=get_font(font_size)).pack(side="right", padx=20)
        
        # Non-Inventory Sales
        non_inv_row = ctk.CTkFrame(sales_frame)
        non_inv_row.pack(fill="x", pady=2)
        ctk.CTkLabel(non_inv_row, text="🍽️ Non-Inventory Sales", font=get_font(font_size), width=200).pack(side="left", padx=20)
        ctk.CTkLabel(non_inv_row, text=f"UGX {sales_summary['non_inventory']['revenue']:,.0f} ({sales_summary['non_inventory']['count']} items)", font=get_font(font_size)).pack(side="right", padx=20)
        
        # Room Sales
        room_row = ctk.CTkFrame(sales_frame)
        room_row.pack(fill="x", pady=2)
        ctk.CTkLabel(room_row, text="🛏️ Room Sales", font=get_font(font_size), width=200).pack(side="left", padx=20)
        ctk.CTkLabel(room_row, text=f"UGX {sales_summary['room']['revenue']:,.0f} ({sales_summary['room']['count']} bookings)", font=get_font(font_size)).pack(side="right", padx=20)
        
        # Sales Details
        if sales_summary['details']['inventory'] or sales_summary['details']['non_inventory'] or sales_summary['details']['room']:
            detail_frame = ctk.CTkFrame(self.report_frame, fg_color="#1a2a3a")
            detail_frame.pack(fill="x", pady=5)
            
            ctk.CTkLabel(detail_frame, text="SALES DETAILS", font=get_font_bold(font_size + 2)).pack(pady=5)
            
            # Inventory Sales
            if sales_summary['details']['inventory']:
                ctk.CTkLabel(detail_frame, text="📦 Inventory Items Sold", font=get_font_bold(font_size)).pack(anchor="w", padx=20, pady=2)
                for sale in sales_summary['details']['inventory'][:10]:
                    row = ctk.CTkFrame(detail_frame)
                    row.pack(fill="x", pady=1)
                    ctk.CTkLabel(row, text=f"{sale[1]} x{sale[2]}", font=get_font(font_size - 1), width=200).pack(side="left", padx=20)
                    ctk.CTkLabel(row, text=f"UGX {sale[4]:,.0f}", font=get_font(font_size - 1)).pack(side="right", padx=20)
            
            # Non-Inventory Sales
            if sales_summary['details']['non_inventory']:
                ctk.CTkLabel(detail_frame, text="🍽️ Non-Inventory Items Sold", font=get_font_bold(font_size)).pack(anchor="w", padx=20, pady=2)
                for sale in sales_summary['details']['non_inventory'][:10]:
                    row = ctk.CTkFrame(detail_frame)
                    row.pack(fill="x", pady=1)
                    ctk.CTkLabel(row, text=f"{sale[1]} x{sale[2]}", font=get_font(font_size - 1), width=200).pack(side="left", padx=20)
                    ctk.CTkLabel(row, text=f"UGX {sale[4]:,.0f}", font=get_font(font_size - 1)).pack(side="right", padx=20)
            
            # Room Sales
            if sales_summary['details']['room']:
                ctk.CTkLabel(detail_frame, text="🛏️ Room Bookings Checked Out", font=get_font_bold(font_size)).pack(anchor="w", padx=20, pady=2)
                for sale in sales_summary['details']['room'][:10]:
                    row = ctk.CTkFrame(detail_frame)
                    row.pack(fill="x", pady=1)
                    ctk.CTkLabel(row, text=f"{sale[1]}", font=get_font(font_size - 1), width=200).pack(side="left", padx=20)
                    ctk.CTkLabel(row, text=f"UGX {sale[4]:,.0f}", font=get_font(font_size - 1)).pack(side="right", padx=20)
        
        # Expenses Section - Detailed
        expenses_frame = ctk.CTkFrame(self.report_frame, fg_color="#1a2a3a")
        expenses_frame.pack(fill="x", pady=5)
        
        ctk.CTkLabel(expenses_frame, text="EXPENSES DETAIL", font=get_font_bold(font_size + 2)).pack(pady=5)
        
        # Summary row
        summary_row = ctk.CTkFrame(expenses_frame)
        summary_row.pack(fill="x", pady=5)
        
        ctk.CTkLabel(
            summary_row,
            text=f"Total: UGX {expense_summary['total']:,.0f} | Count: {expense_summary['count']} | Avg: UGX {expense_summary['average']:,.0f} | Min: UGX {expense_summary['min']:,.0f} | Max: UGX {expense_summary['max']:,.0f}",
            font=get_font(font_size)
        ).pack(pady=5)
        
        if expenses:
            # Header
            header = ctk.CTkFrame(expenses_frame)
            header.pack(fill="x", pady=2)
            
            headers = ["#", "Name", "Amount", "Time"]
            widths = [30, 250, 120, 150]
            
            for h, w in zip(headers, widths):
                ctk.CTkLabel(
                    header,
                    text=h,
                    font=get_font_bold(font_size),
                    width=w
                ).pack(side="left", padx=5)
            
            for idx, expense in enumerate(expenses[:20], 1):
                exp_id, name, amount, date_str, notes = expense
                
                row = ctk.CTkFrame(expenses_frame)
                row.pack(fill="x", pady=1)
                
                data = [
                    str(idx),
                    name,
                    f"UGX {amount:,.0f}",
                    date_str[:16] if date_str else ""
                ]
                
                for i, d in enumerate(data):
                    ctk.CTkLabel(
                        row,
                        text=d,
                        width=widths[i],
                        anchor="w" if i > 0 else "center",
                        font=get_font(font_size)
                    ).pack(side="left", padx=5)
            
            if len(expenses) > 20:
                ctk.CTkLabel(
                    expenses_frame,
                    text=f"... and {len(expenses) - 20} more expenses",
                    font=get_font(font_size - 2),
                    text_color="gray"
                ).pack(pady=2)
        else:
            ctk.CTkLabel(
                expenses_frame,
                text="No expenses recorded on this day",
                font=get_font(font_size),
                text_color="gray"
            ).pack(pady=5)
        
        # Net Profit
        profit_frame = ctk.CTkFrame(self.report_frame, fg_color="#1a3a1a" if report['net_profit'] >= 0 else "#3a1a1a")
        profit_frame.pack(fill="x", pady=5)
        
        ctk.CTkLabel(profit_frame, text="NET PROFIT", font=get_font_bold(font_size + 2)).pack(pady=5)
        
        profit_row = ctk.CTkFrame(profit_frame)
        profit_row.pack(fill="x", pady=2)
        ctk.CTkLabel(profit_row, text="Net Profit", font=get_font(font_size), width=200).pack(side="left", padx=20)
        ctk.CTkLabel(
            profit_row, 
            text=f"UGX {report['net_profit']:,.0f}", 
            font=get_font_bold(font_size + 4),
            text_color="green" if report['net_profit'] >= 0 else "red"
        ).pack(side="right", padx=20)
        
        # Rooms
        rooms_frame = ctk.CTkFrame(self.report_frame, fg_color="#1a2a3a")
        rooms_frame.pack(fill="x", pady=5)
        
        ctk.CTkLabel(rooms_frame, text="ROOMS", font=get_font_bold(font_size + 2)).pack(pady=5)
        
        room_row = ctk.CTkFrame(rooms_frame)
        room_row.pack(fill="x", pady=2)
        ctk.CTkLabel(room_row, text="Occupied Rooms", font=get_font(font_size), width=200).pack(side="left", padx=20)
        ctk.CTkLabel(room_row, text=f"{report['occupied']}/{report['total_rooms']}", font=get_font_bold(font_size)).pack(side="right", padx=20)
        
        occupancy = (report['occupied'] / report['total_rooms'] * 100) if report['total_rooms'] > 0 else 0
        room_row2 = ctk.CTkFrame(rooms_frame)
        room_row2.pack(fill="x", pady=2)
        ctk.CTkLabel(room_row2, text="Occupancy Rate", font=get_font(font_size), width=200).pack(side="left", padx=20)
        ctk.CTkLabel(room_row2, text=f"{occupancy:.1f}%", font=get_font_bold(font_size)).pack(side="right", padx=20)
        
        # Store for export
        self.current_report = report
        self.current_expenses = expenses
        self.current_expense_summary = expense_summary
        self.current_sales_summary = sales_summary
        self.current_report_date = date_str
        self.report_type = "daily"
    
    def show_daily_report(self):
        """Show daily report for today"""
        today = datetime.now().strftime("%Y-%m-%d")
        self.date_entry.delete(0, "end")
        self.date_entry.insert(0, today)
        self.show_daily_report_for_date(today)
    
    def clear_report(self):
        """Clear the report frame"""
        for widget in self.report_frame.winfo_children():
            widget.destroy()
    
    def get_sales_with_types(self, start_date=None, end_date=None):
        """Get sales with inventory/non-inventory classification"""
        conn = get_connection()
        cursor = conn.cursor()
        
        if start_date and end_date:
            cursor.execute("""
                SELECT id, item_name, quantity, unit_price, total_revenue, 
                       total_cost, total_profit, sale_date, notes, payment_type
                FROM sales
                WHERE DATE(sale_date) >= ? AND DATE(sale_date) <= ?
                ORDER BY sale_date DESC
            """, (start_date, end_date))
        elif start_date:
            cursor.execute("""
                SELECT id, item_name, quantity, unit_price, total_revenue, 
                       total_cost, total_profit, sale_date, notes, payment_type
                FROM sales
                WHERE DATE(sale_date) >= ?
                ORDER BY sale_date DESC
            """, (start_date,))
        else:
            cursor.execute("""
                SELECT id, item_name, quantity, unit_price, total_revenue, 
                       total_cost, total_profit, sale_date, notes, payment_type
                FROM sales
                ORDER BY sale_date DESC
            """)
        
        sales = cursor.fetchall()
        conn.close()
        
        # Classify each sale
        inventory_sales = []
        non_inventory_sales = []
        room_sales = []
        
        for sale in sales:
            item_name = sale[1]
            # Check if it's a room sale
            if item_name.startswith("Room:"):
                room_sales.append(sale)
            elif get_item_type(item_name) == "inventory":
                inventory_sales.append(sale)
            else:
                non_inventory_sales.append(sale)
        
        return {
            "all_sales": sales,
            "inventory": inventory_sales,
            "non_inventory": non_inventory_sales,
            "room": room_sales
        }
    
    def get_sales_summary_with_types(self, start_date=None, end_date=None):
        """Get sales summary broken down by type"""
        sales_data = self.get_sales_with_types(start_date, end_date)
        
        def calc_summary(sales_list):
            if not sales_list:
                return {"count": 0, "revenue": 0, "profit": 0, "cost": 0}
            
            count = len(sales_list)
            revenue = sum(s[4] for s in sales_list)  # total_revenue
            profit = sum(s[6] for s in sales_list)   # total_profit
            cost = sum(s[5] for s in sales_list)     # total_cost
            return {"count": count, "revenue": revenue, "profit": profit, "cost": cost}
        
        return {
            "total": calc_summary(sales_data["all_sales"]),
            "inventory": calc_summary(sales_data["inventory"]),
            "non_inventory": calc_summary(sales_data["non_inventory"]),
            "room": calc_summary(sales_data["room"]),
            "details": {
                "inventory": sales_data["inventory"],
                "non_inventory": sales_data["non_inventory"],
                "room": sales_data["room"]
            }
        }
    
    def get_expenses_for_period(self, start_date=None, end_date=None):
        """Get expenses for a specific period with details"""
        conn = get_connection()
        cursor = conn.cursor()
        
        if start_date and end_date:
            cursor.execute("""
                SELECT id, name, amount, expense_date, notes
                FROM expenses
                WHERE DATE(expense_date) >= ? AND DATE(expense_date) <= ?
                ORDER BY expense_date DESC
            """, (start_date, end_date))
        elif start_date:
            cursor.execute("""
                SELECT id, name, amount, expense_date, notes
                FROM expenses
                WHERE DATE(expense_date) >= ?
                ORDER BY expense_date DESC
            """, (start_date,))
        else:
            cursor.execute("""
                SELECT id, name, amount, expense_date, notes
                FROM expenses
                ORDER BY expense_date DESC
            """)
        
        expenses = cursor.fetchall()
        conn.close()
        return expenses
    
    def get_expenses_summary(self, start_date=None, end_date=None):
        """Get expense summary for a period"""
        conn = get_connection()
        cursor = conn.cursor()
        
        if start_date and end_date:
            cursor.execute("""
                SELECT 
                    COUNT(*) as count,
                    SUM(amount) as total,
                    AVG(amount) as average,
                    MIN(amount) as min_amount,
                    MAX(amount) as max_amount
                FROM expenses
                WHERE DATE(expense_date) >= ? AND DATE(expense_date) <= ?
            """, (start_date, end_date))
        elif start_date:
            cursor.execute("""
                SELECT 
                    COUNT(*) as count,
                    SUM(amount) as total,
                    AVG(amount) as average,
                    MIN(amount) as min_amount,
                    MAX(amount) as max_amount
                FROM expenses
                WHERE DATE(expense_date) >= ?
            """, (start_date,))
        else:
            cursor.execute("""
                SELECT 
                    COUNT(*) as count,
                    SUM(amount) as total,
                    AVG(amount) as average,
                    MIN(amount) as min_amount,
                    MAX(amount) as max_amount
                FROM expenses
            """)
        
        result = cursor.fetchone()
        conn.close()
        
        if result:
            return {
                "count": result[0] or 0,
                "total": result[1] or 0,
                "average": result[2] or 0,
                "min": result[3] or 0,
                "max": result[4] or 0
            }
        return {"count": 0, "total": 0, "average": 0, "min": 0, "max": 0}
    
    def show_weekly_report(self):
        """Show weekly report with profit - Enhanced with more details"""
        self.clear_report()
        
        font_size = get_font_size()
        today = datetime.now().strftime("%Y-%m-%d")
        
        # Get weekly data
        sales_data, expenses_data = get_weekly_report()
        
        # Get business date
        business_date = get_business_date()
        week_start = get_weekly_start_date(business_date)
        week_end = business_date
        
        # Get detailed expenses for the week
        week_expenses_detail = self.get_expenses_for_period(week_start, week_end)
        week_expense_summary = self.get_expenses_summary(week_start, week_end)
        
        # Get sales summary with types for the week
        week_sales_summary = self.get_sales_summary_with_types(week_start, week_end)
        
        # Title
        ctk.CTkLabel(
            self.report_frame,
            text="VALLEY VIEW MOTEL",
            font=get_font_bold(24),
            text_color="#1f538d"
        ).pack(pady=5)
        
        ctk.CTkLabel(
            self.report_frame,
            text="📈 Weekly Performance Report",
            font=get_font_bold(20)
        ).pack(pady=5)
        
        week_start_display = datetime.strptime(week_start, "%Y-%m-%d").strftime("%b %d, %Y")
        week_end_display = datetime.strptime(week_end, "%Y-%m-%d").strftime("%b %d, %Y")
        
        ctk.CTkLabel(
            self.report_frame,
            text=f"Week: {week_start_display} - {week_end_display}",
            font=get_font(font_size + 2),
            text_color="gray"
        ).pack(pady=2)
        
        ctk.CTkLabel(
            self.report_frame,
            text=f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            font=get_font(font_size - 2),
            text_color="gray"
        ).pack(pady=2)
        
        ctk.CTkFrame(self.report_frame, height=2, fg_color="gray").pack(fill="x", pady=10)
        
        # Combine data
        weekly_data = {}
        for date, revenue, cost, profit in sales_data:
            weekly_data[date] = {
                "revenue": revenue or 0,
                "cost": cost or 0,
                "profit": profit or 0,
                "expenses": 0
            }
        for date, expenses in expenses_data:
            if date in weekly_data:
                weekly_data[date]["expenses"] = expenses or 0
            else:
                weekly_data[date] = {
                    "revenue": 0,
                    "cost": 0,
                    "profit": 0,
                    "expenses": expenses or 0
                }
        
        if not weekly_data:
            ctk.CTkLabel(
                self.report_frame,
                text="No data for the past week",
                font=get_font(font_size + 2),
                text_color="gray"
            ).pack(pady=20)
            return
        
        # Summary Cards at top
        summary_frame = ctk.CTkFrame(self.report_frame, fg_color="#1a2a3a")
        summary_frame.pack(fill="x", pady=5)
        
        ctk.CTkLabel(summary_frame, text="WEEKLY SUMMARY", font=get_font_bold(font_size + 2)).pack(pady=5)
        
        total_revenue = 0
        total_cost = 0
        total_gross_profit = 0
        total_expenses = 0
        
        for date, data in weekly_data.items():
            total_revenue += data["revenue"]
            total_cost += data["cost"]
            total_gross_profit += data["profit"]
            total_expenses += data["expenses"]
        
        total_net_profit = total_gross_profit - total_expenses
        
        summary_stats = [
            ("Total Revenue", f"UGX {total_revenue:,.0f}"),
            ("Total Cost of Goods", f"UGX {total_cost:,.0f}"),
            ("Gross Profit", f"UGX {total_gross_profit:,.0f}"),
            ("Total Expenses", f"UGX {total_expenses:,.0f}"),
            ("Net Profit", f"UGX {total_net_profit:,.0f}")
        ]
        
        for label, value in summary_stats:
            row = ctk.CTkFrame(summary_frame)
            row.pack(fill="x", pady=2)
            ctk.CTkLabel(row, text=label, font=get_font(font_size), width=200).pack(side="left", padx=20)
            ctk.CTkLabel(row, text=value, font=get_font_bold(font_size)).pack(side="right", padx=20)
        
        # Sales Type Breakdown for the Week
        sales_type_frame = ctk.CTkFrame(self.report_frame, fg_color="#1a2a3a")
        sales_type_frame.pack(fill="x", pady=5)
        
        ctk.CTkLabel(sales_type_frame, text="SALES BY TYPE", font=get_font_bold(font_size + 2)).pack(pady=5)
        
        type_stats = [
            ("📦 Inventory", f"UGX {week_sales_summary['inventory']['revenue']:,.0f}", f"{week_sales_summary['inventory']['count']} items"),
            ("🍽️ Non-Inventory", f"UGX {week_sales_summary['non_inventory']['revenue']:,.0f}", f"{week_sales_summary['non_inventory']['count']} items"),
            ("🛏️ Rooms", f"UGX {week_sales_summary['room']['revenue']:,.0f}", f"{week_sales_summary['room']['count']} bookings")
        ]
        
        for label, value, sub in type_stats:
            row = ctk.CTkFrame(sales_type_frame)
            row.pack(fill="x", pady=2)
            ctk.CTkLabel(row, text=label, font=get_font(font_size), width=180).pack(side="left", padx=20)
            ctk.CTkLabel(row, text=value, font=get_font_bold(font_size)).pack(side="right", padx=20)
            ctk.CTkLabel(row, text=sub, font=get_font(font_size - 2), text_color="gray").pack(side="right", padx=5)
        
        # Daily Breakdown
        ctk.CTkLabel(
            self.report_frame,
            text="📊 DAILY BREAKDOWN",
            font=get_font_bold(font_size + 2)
        ).pack(pady=10)
        
        # Header
        header = ctk.CTkFrame(self.report_frame)
        header.pack(fill="x", pady=5)
        
        headers = ["Date", "Day", "Revenue", "Cost", "Gross Profit", "Expenses", "Net Profit"]
        widths = [100, 60, 120, 120, 120, 120, 120]
        
        for h, w in zip(headers, widths):
            ctk.CTkLabel(
                header,
                text=h,
                font=get_font_bold(font_size),
                width=w
            ).pack(side="left", padx=5)
        
        for date in sorted(weekly_data.keys()):
            data = weekly_data[date]
            revenue = data["revenue"]
            cost = data["cost"]
            gross_profit = data["profit"]
            expenses = data["expenses"]
            net_profit = gross_profit - expenses
            
            # Get day name
            date_obj = datetime.strptime(date, "%Y-%m-%d")
            day_name = ["Mon", "Tue", "Wed", "Thu", "Fri", "Sat", "Sun"][date_obj.weekday()]
            
            row = ctk.CTkFrame(self.report_frame)
            
            # Highlight Saturday
            if day_name == "Sat":
                row.configure(fg_color="#1a3a1a")
            
            row.pack(fill="x", pady=2)
            
            row_data = [
                date,
                day_name,
                f"UGX {revenue:,.0f}",
                f"UGX {cost:,.0f}",
                f"UGX {gross_profit:,.0f}",
                f"UGX {expenses:,.0f}",
                f"UGX {net_profit:,.0f}"
            ]
            
            for i, d in enumerate(row_data):
                if i == 6:  # Net Profit column
                    color = "green" if net_profit >= 0 else "red"
                    ctk.CTkLabel(
                        row,
                        text=d,
                        width=widths[i],
                        text_color=color,
                        font=get_font_bold(font_size)
                    ).pack(side="left", padx=5)
                else:
                    ctk.CTkLabel(
                        row,
                        text=d,
                        width=widths[i],
                        font=get_font(font_size)
                    ).pack(side="left", padx=5)
        
        # Totals
        total_frame = ctk.CTkFrame(self.report_frame, fg_color="#1a2a3a")
        total_frame.pack(fill="x", pady=10)
        
        ctk.CTkLabel(
            total_frame,
            text=f"Total Revenue: UGX {total_revenue:,.0f}",
            font=get_font_bold(font_size + 2)
        ).pack(side="left", padx=20)
        
        ctk.CTkLabel(
            total_frame,
            text=f"Net Profit: UGX {total_net_profit:,.0f}",
            font=get_font_bold(font_size + 2),
            text_color="green" if total_net_profit >= 0 else "red"
        ).pack(side="left", padx=20)
        
        # Weekly Expenses Detail
        expense_frame = ctk.CTkFrame(self.report_frame, fg_color="#1a2a3a")
        expense_frame.pack(fill="x", pady=10)
        
        ctk.CTkLabel(
            expense_frame,
            text="📝 WEEKLY EXPENSES DETAIL",
            font=get_font_bold(font_size + 2)
        ).pack(pady=5)
        
        # Expense summary
        exp_summary_row = ctk.CTkFrame(expense_frame)
        exp_summary_row.pack(fill="x", pady=5)
        
        ctk.CTkLabel(
            exp_summary_row,
            text=f"Total: UGX {week_expense_summary['total']:,.0f} | Count: {week_expense_summary['count']} | Avg: UGX {week_expense_summary['average']:,.0f}",
            font=get_font(font_size)
        ).pack(pady=2)
        
        if week_expenses_detail:
            # Header
            exp_header = ctk.CTkFrame(expense_frame)
            exp_header.pack(fill="x", pady=2)
            
            exp_headers = ["#", "Name", "Amount", "Date", "Notes"]
            exp_widths = [30, 180, 120, 120, 150]
            
            for h, w in zip(exp_headers, exp_widths):
                ctk.CTkLabel(
                    exp_header,
                    text=h,
                    font=get_font_bold(font_size),
                    width=w
                ).pack(side="left", padx=5)
            
            for idx, expense in enumerate(week_expenses_detail[:15], 1):
                exp_id, name, amount, date_str, notes = expense
                
                row = ctk.CTkFrame(expense_frame)
                row.pack(fill="x", pady=1)
                
                data = [
                    str(idx),
                    name,
                    f"UGX {amount:,.0f}",
                    date_str[:10] if date_str else "",
                    notes[:20] + "..." if notes and len(notes) > 20 else notes or ""
                ]
                
                for i, d in enumerate(data):
                    ctk.CTkLabel(
                        row,
                        text=d,
                        width=exp_widths[i],
                        anchor="w" if i > 0 else "center",
                        font=get_font(font_size - 1)
                    ).pack(side="left", padx=5)
            
            if len(week_expenses_detail) > 15:
                ctk.CTkLabel(
                    expense_frame,
                    text=f"... and {len(week_expenses_detail) - 15} more expenses",
                    font=get_font(font_size - 2),
                    text_color="gray"
                ).pack(pady=2)
        else:
            ctk.CTkLabel(
                expense_frame,
                text="No expenses recorded this week",
                font=get_font(font_size),
                text_color="gray"
            ).pack(pady=5)
        
        # Store for export
        self.current_report = {
            "date": "Weekly",
            "week_start": week_start,
            "week_end": week_end,
            "revenue": total_revenue,
            "cost": total_cost,
            "gross_profit": total_gross_profit,
            "expenses": total_expenses,
            "net_profit": total_net_profit,
            "weekly_data": weekly_data,
            "expenses_detail": week_expenses_detail,
            "expense_summary": week_expense_summary,
            "sales_summary": week_sales_summary
        }
        self.report_type = "weekly"
    
    def show_balance_report(self):
        """Show weekly balance report with opening/closing balances"""
        self.clear_report()
        
        font_size = get_font_size()
        today = datetime.now().strftime("%Y-%m-%d")
        weekly_data = get_weekly_report_data(today)
        
        # Title
        ctk.CTkLabel(
            self.report_frame,
            text="VALLEY VIEW MOTEL",
            font=get_font_bold(24),
            text_color="#1f538d"
        ).pack(pady=5)
        
        ctk.CTkLabel(
            self.report_frame,
            text="📊 Weekly Balance Report",
            font=get_font_bold(20)
        ).pack(pady=5)
        
        ctk.CTkLabel(
            self.report_frame,
            text=f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            font=get_font(font_size - 2),
            text_color="gray"
        ).pack(pady=2)
        
        ctk.CTkFrame(self.report_frame, height=2, fg_color="gray").pack(fill="x", pady=10)
        
        if not weekly_data:
            ctk.CTkLabel(
                self.report_frame,
                text="No data for this week",
                font=get_font(font_size + 2),
                text_color="gray"
            ).pack(pady=20)
            return
        
        # Header
        header = ctk.CTkFrame(self.report_frame)
        header.pack(fill="x", pady=5)
        
        headers = ["Date", "Opening", "Revenue", "Expenses", "Closing", "Cash Collected", "Merchant", "Day"]
        widths = [100, 100, 100, 100, 100, 120, 120, 80]
        
        for h, w in zip(headers, widths):
            ctk.CTkLabel(
                header,
                text=h,
                font=get_font_bold(font_size),
                width=w
            ).pack(side="left", padx=5)
        
        total_revenue = 0
        total_expenses = 0
        total_cash = 0
        total_merchant = 0
        
        for data in weekly_data:
            row = ctk.CTkFrame(self.report_frame)
            row.pack(fill="x", pady=2)
            
            # Highlight Saturday
            if data["is_saturday"]:
                row.configure(fg_color="#1a3a1a")
            
            # Format day name
            date_obj = datetime.strptime(data["date"], "%Y-%m-%d")
            day_name = ["Mon", "Tue", "Wed", "Thu", "Fri", "Sat", "Sun"][date_obj.weekday()]
            
            row_data = [
                data["date"],
                f"UGX {data['opening_balance']:,.0f}",
                f"UGX {data['daily_revenue']:,.0f}",
                f"UGX {data['daily_expenses']:,.0f}",
                f"UGX {data['closing_balance']:,.0f}",
                f"UGX {data['cash_collected']:,.0f}",
                f"UGX {data['merchant_collected']:,.0f}",
                day_name
            ]
            
            if data["is_saturday"]:
                row_data[7] = "📅 SAT (Collection)"
            
            total_revenue += data["daily_revenue"]
            total_expenses += data["daily_expenses"]
            total_cash += data["cash_collected"]
            total_merchant += data["merchant_collected"]
            
            for i, d in enumerate(row_data):
                if i == 4:  # Closing balance
                    color = "green" if data["closing_balance"] >= 0 else "red"
                    ctk.CTkLabel(
                        row,
                        text=d,
                        width=widths[i],
                        text_color=color,
                        font=get_font_bold(font_size)
                    ).pack(side="left", padx=5)
                elif i == 7:  # Day
                    ctk.CTkLabel(
                        row,
                        text=d,
                        width=widths[i],
                        font=get_font_bold(font_size) if data["is_saturday"] else get_font(font_size)
                    ).pack(side="left", padx=5)
                else:
                    ctk.CTkLabel(
                        row,
                        text=d,
                        width=widths[i],
                        anchor="w" if i > 0 else "center",
                        font=get_font(font_size)
                    ).pack(side="left", padx=5)
        
        # Totals
        total_frame = ctk.CTkFrame(self.report_frame, fg_color="#1a2a3a")
        total_frame.pack(fill="x", pady=10)
        
        ctk.CTkLabel(
            total_frame,
            text=f"Total Revenue: UGX {total_revenue:,.0f}",
            font=get_font_bold(font_size + 2)
        ).pack(side="left", padx=20)
        
        ctk.CTkLabel(
            total_frame,
            text=f"Total Expenses: UGX {total_expenses:,.0f}",
            font=get_font_bold(font_size + 2)
        ).pack(side="left", padx=20)
        
        ctk.CTkLabel(
            total_frame,
            text=f"Total Cash Collected: UGX {total_cash:,.0f}",
            font=get_font_bold(font_size + 2),
            text_color="green"
        ).pack(side="left", padx=20)
        
        ctk.CTkLabel(
            total_frame,
            text=f"Total Merchant: UGX {total_merchant:,.0f}",
            font=get_font_bold(font_size + 2),
            text_color="blue"
        ).pack(side="left", padx=20)
        
        # Net Balance
        net_balance = total_revenue - total_expenses
        net_frame = ctk.CTkFrame(self.report_frame, fg_color="#1a3a1a" if net_balance >= 0 else "#3a1a1a")
        net_frame.pack(fill="x", pady=10)
        
        ctk.CTkLabel(
            net_frame,
            text="NET BALANCE FOR WEEK",
            font=get_font_bold(font_size + 2)
        ).pack(side="left", padx=20)
        
        ctk.CTkLabel(
            net_frame,
            text=f"UGX {net_balance:,.0f}",
            font=get_font_bold(font_size + 4),
            text_color="green" if net_balance >= 0 else "red"
        ).pack(side="right", padx=20)
        
        # Store for export
        self.current_report = {
            "date": "Balance Report",
            "weekly_data": weekly_data,
            "total_revenue": total_revenue,
            "total_expenses": total_expenses,
            "total_cash": total_cash,
            "total_merchant": total_merchant,
            "net_balance": net_balance
        }
        self.report_type = "balance"
    
    def show_all_data_report(self):
        """Show all data report - all sales and expenses ever recorded"""
        self.clear_report()
        
        font_size = get_font_size()
        
        # Title
        ctk.CTkLabel(
            self.report_frame,
            text="VALLEY VIEW MOTEL",
            font=get_font_bold(24),
            text_color="#1f538d"
        ).pack(pady=5)
        
        ctk.CTkLabel(
            self.report_frame,
            text="📊 Lifetime Data Report",
            font=get_font_bold(20)
        ).pack(pady=5)
        
        ctk.CTkLabel(
            self.report_frame,
            text=f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            font=get_font(font_size - 2),
            text_color="gray"
        ).pack(pady=2)
        
        ctk.CTkFrame(self.report_frame, height=2, fg_color="gray").pack(fill="x", pady=10)
        
        # Get all sales
        try:
            conn = get_connection()
            cursor = conn.cursor()
            
            # Get total sales
            cursor.execute("""
                SELECT COUNT(*), SUM(total_revenue), SUM(total_profit)
                FROM sales
            """)
            sales_total = cursor.fetchone()
            
            # Get total expenses
            cursor.execute("""
                SELECT COUNT(*), SUM(amount)
                FROM expenses
            """)
            expenses_total = cursor.fetchone()
            
            # Get sales by payment type
            cursor.execute("""
                SELECT payment_type, SUM(total_revenue), COUNT(*)
                FROM sales
                GROUP BY payment_type
            """)
            payment_types = cursor.fetchall()
            
            # Get monthly breakdown
            cursor.execute("""
                SELECT strftime('%Y-%m', sale_date) as month, 
                       SUM(total_revenue), 
                       SUM(total_profit),
                       COUNT(*)
                FROM sales
                GROUP BY strftime('%Y-%m', sale_date)
                ORDER BY month DESC
            """)
            monthly_data = cursor.fetchall()
            
            # Get all expenses for detail
            cursor.execute("""
                SELECT id, name, amount, expense_date, notes
                FROM expenses
                ORDER BY expense_date DESC
                LIMIT 30
            """)
            all_expenses = cursor.fetchall()
            
            # Get lifetime sales by type
            lifetime_sales_summary = self.get_sales_summary_with_types()
            
            conn.close()
        except Exception as e:
            print(f"Error getting all data: {e}")
            sales_total = (0, 0, 0)
            expenses_total = (0, 0)
            payment_types = []
            monthly_data = []
            all_expenses = []
            lifetime_sales_summary = {"inventory": {"count": 0, "revenue": 0}, "non_inventory": {"count": 0, "revenue": 0}, "room": {"count": 0, "revenue": 0}}
        
        # Summary Cards
        summary_frame = ctk.CTkFrame(self.report_frame, fg_color="#1a2a3a")
        summary_frame.pack(fill="x", pady=5)
        
        ctk.CTkLabel(
            summary_frame,
            text="📊 LIFETIME TOTALS",
            font=get_font_bold(font_size + 2)
        ).pack(pady=5)
        
        total_sales_count, total_revenue, total_profit = sales_total
        total_exp_count, total_exp_amount = expenses_total
        
        summary_stats = [
            ("Total Sales", str(total_sales_count or 0)),
            ("Total Revenue", f"UGX {(total_revenue or 0):,.0f}"),
            ("Total Profit", f"UGX {(total_profit or 0):,.0f}"),
            ("Total Expenses", f"UGX {(total_exp_amount or 0):,.0f}"),
            ("Net Profit", f"UGX {((total_profit or 0) - (total_exp_amount or 0)):,.0f}")
        ]
        
        for label, value in summary_stats:
            row = ctk.CTkFrame(summary_frame)
            row.pack(fill="x", pady=2)
            ctk.CTkLabel(row, text=label, font=get_font(font_size), width=200).pack(side="left", padx=20)
            ctk.CTkLabel(row, text=value, font=get_font_bold(font_size)).pack(side="right", padx=20)
        
        # Sales by Type (Lifetime)
        type_frame = ctk.CTkFrame(self.report_frame, fg_color="#1a2a3a")
        type_frame.pack(fill="x", pady=5)
        
        ctk.CTkLabel(type_frame, text="SALES BY TYPE (LIFETIME)", font=get_font_bold(font_size + 2)).pack(pady=5)
        
        type_stats = [
            ("📦 Inventory", f"UGX {lifetime_sales_summary['inventory']['revenue']:,.0f}", f"{lifetime_sales_summary['inventory']['count']} items"),
            ("🍽️ Non-Inventory", f"UGX {lifetime_sales_summary['non_inventory']['revenue']:,.0f}", f"{lifetime_sales_summary['non_inventory']['count']} items"),
            ("🛏️ Rooms", f"UGX {lifetime_sales_summary['room']['revenue']:,.0f}", f"{lifetime_sales_summary['room']['count']} bookings")
        ]
        
        for label, value, sub in type_stats:
            row = ctk.CTkFrame(type_frame)
            row.pack(fill="x", pady=2)
            ctk.CTkLabel(row, text=label, font=get_font(font_size), width=180).pack(side="left", padx=20)
            ctk.CTkLabel(row, text=value, font=get_font_bold(font_size)).pack(side="right", padx=20)
            ctk.CTkLabel(row, text=sub, font=get_font(font_size - 2), text_color="gray").pack(side="right", padx=5)
        
        # Payment Type Breakdown
        if payment_types:
            payment_frame = ctk.CTkFrame(self.report_frame, fg_color="#1a2a3a")
            payment_frame.pack(fill="x", pady=5)
            
            ctk.CTkLabel(
                payment_frame,
                text="PAYMENT TYPE BREAKDOWN",
                font=get_font_bold(font_size + 2)
            ).pack(pady=5)
            
            for ptype, total, count in payment_types:
                row = ctk.CTkFrame(payment_frame)
                row.pack(fill="x", pady=2)
                ctk.CTkLabel(row, text=f"{ptype}", font=get_font(font_size), width=200).pack(side="left", padx=20)
                ctk.CTkLabel(row, text=f"UGX {total:,.0f} ({count} transactions)", font=get_font_bold(font_size)).pack(side="right", padx=20)
        
        # Monthly Breakdown
        if monthly_data:
            monthly_frame = ctk.CTkFrame(self.report_frame, fg_color="#1a2a3a")
            monthly_frame.pack(fill="x", pady=5)
            
            ctk.CTkLabel(
                monthly_frame,
                text="📅 MONTHLY BREAKDOWN",
                font=get_font_bold(font_size + 2)
            ).pack(pady=5)
            
            for month, revenue, profit, count in monthly_data:
                row = ctk.CTkFrame(monthly_frame)
                row.pack(fill="x", pady=2)
                ctk.CTkLabel(row, text=month, font=get_font(font_size), width=150).pack(side="left", padx=20)
                ctk.CTkLabel(row, text=f"UGX {revenue:,.0f} | Profit: UGX {profit:,.0f} ({count} sales)", font=get_font(font_size)).pack(side="right", padx=20)
        
        # All Expenses Detail
        expense_frame = ctk.CTkFrame(self.report_frame, fg_color="#1a2a3a")
        expense_frame.pack(fill="x", pady=10)
        
        ctk.CTkLabel(
            expense_frame,
            text="📝 RECENT EXPENSES (Last 30)",
            font=get_font_bold(font_size + 2)
        ).pack(pady=5)
        
        if all_expenses:
            # Header
            exp_header = ctk.CTkFrame(expense_frame)
            exp_header.pack(fill="x", pady=2)
            
            exp_headers = ["#", "Name", "Amount", "Date", "Notes"]
            exp_widths = [30, 180, 120, 120, 150]
            
            for h, w in zip(exp_headers, exp_widths):
                ctk.CTkLabel(
                    exp_header,
                    text=h,
                    font=get_font_bold(font_size),
                    width=w
                ).pack(side="left", padx=5)
            
            for idx, expense in enumerate(all_expenses[:20], 1):
                exp_id, name, amount, date_str, notes = expense
                
                row = ctk.CTkFrame(expense_frame)
                row.pack(fill="x", pady=1)
                
                data = [
                    str(idx),
                    name,
                    f"UGX {amount:,.0f}",
                    date_str[:10] if date_str else "",
                    notes[:20] + "..." if notes and len(notes) > 20 else notes or ""
                ]
                
                for i, d in enumerate(data):
                    ctk.CTkLabel(
                        row,
                        text=d,
                        width=exp_widths[i],
                        anchor="w" if i > 0 else "center",
                        font=get_font(font_size - 1)
                    ).pack(side="left", padx=5)
        else:
            ctk.CTkLabel(
                expense_frame,
                text="No expenses recorded",
                font=get_font(font_size),
                text_color="gray"
            ).pack(pady=5)
        
        # Store for export
        self.current_report = {
            "date": "All Data",
            "total_sales": total_sales_count or 0,
            "total_revenue": total_revenue or 0,
            "total_profit": total_profit or 0,
            "total_expenses": total_exp_amount or 0,
            "net_profit": (total_profit or 0) - (total_exp_amount or 0),
            "payment_types": payment_types,
            "monthly_data": monthly_data,
            "expenses_detail": all_expenses,
            "sales_summary": lifetime_sales_summary
        }
        self.report_type = "all_data"
    
    def export_csv(self):
        """Export CSV report"""
        if not hasattr(self, 'current_report'):
            self.show_message("Please generate a report first", "red")
            return
        
        report = self.current_report
        filename = f"Valley_View_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
        
        try:
            with open(filename, 'w', newline='', encoding='utf-8') as f:
                writer = csv.writer(f)
                
                writer.writerow(["VALLEY VIEW MOTEL MANAGEMENT SYSTEM"])
                writer.writerow([f"Report: {report.get('date', 'Unknown')}"])
                writer.writerow([f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"])
                writer.writerow([])
                
                # Write report data based on type
                if report.get('date') == "All Data":
                    writer.writerow(["LIFETIME TOTALS"])
                    writer.writerow(["Metric", "Value"])
                    writer.writerow(["Total Sales", report.get('total_sales', 0)])
                    writer.writerow(["Total Revenue", report.get('total_revenue', 0)])
                    writer.writerow(["Total Profit", report.get('total_profit', 0)])
                    writer.writerow(["Total Expenses", report.get('total_expenses', 0)])
                    writer.writerow(["Net Profit", report.get('net_profit', 0)])
                    writer.writerow([])
                    
                    if report.get('sales_summary'):
                        writer.writerow(["SALES BY TYPE"])
                        writer.writerow(["Type", "Revenue", "Count"])
                        writer.writerow(["Inventory", report['sales_summary']['inventory']['revenue'], report['sales_summary']['inventory']['count']])
                        writer.writerow(["Non-Inventory", report['sales_summary']['non_inventory']['revenue'], report['sales_summary']['non_inventory']['count']])
                        writer.writerow(["Rooms", report['sales_summary']['room']['revenue'], report['sales_summary']['room']['count']])
                        writer.writerow([])
                    
                    if report.get('payment_types'):
                        writer.writerow(["PAYMENT TYPE BREAKDOWN"])
                        writer.writerow(["Payment Type", "Amount", "Transactions"])
                        for ptype, total, count in report['payment_types']:
                            writer.writerow([ptype, total, count])
                    
                    if report.get('monthly_data'):
                        writer.writerow([])
                        writer.writerow(["MONTHLY BREAKDOWN"])
                        writer.writerow(["Month", "Revenue", "Profit", "Sales"])
                        for month, revenue, profit, count in report['monthly_data']:
                            writer.writerow([month, revenue, profit, count])
                    
                    if report.get('expenses_detail'):
                        writer.writerow([])
                        writer.writerow(["RECENT EXPENSES"])
                        writer.writerow(["Name", "Amount", "Date", "Notes"])
                        for exp in report['expenses_detail']:
                            writer.writerow([exp[1], exp[2], exp[3], exp[4] or ""])
                
                elif 'weekly_data' in report and report.get('date') == "Balance Report":
                    writer.writerow(["WEEKLY BALANCE REPORT"])
                    writer.writerow(["Date", "Opening", "Revenue", "Expenses", "Closing", "Cash", "Merchant", "Day"])
                    for data in report['weekly_data']:
                        writer.writerow([
                            data["date"],
                            data["opening_balance"],
                            data["daily_revenue"],
                            data["daily_expenses"],
                            data["closing_balance"],
                            data["cash_collected"],
                            data["merchant_collected"],
                            "SAT (Collection)" if data["is_saturday"] else ""
                        ])
                
                elif report.get('date') == "Weekly":
                    writer.writerow(["WEEKLY PERFORMANCE"])
                    writer.writerow(["Metric", "Value"])
                    writer.writerow(["Total Revenue", report.get('revenue', 0)])
                    writer.writerow(["Cost of Goods", report.get('cost', 0)])
                    writer.writerow(["Gross Profit", report.get('gross_profit', 0)])
                    writer.writerow(["Total Expenses", report.get('expenses', 0)])
                    writer.writerow(["Net Profit", report.get('net_profit', 0)])
                    writer.writerow([])
                    
                    if report.get('sales_summary'):
                        writer.writerow(["SALES BY TYPE"])
                        writer.writerow(["Type", "Revenue", "Count"])
                        writer.writerow(["Inventory", report['sales_summary']['inventory']['revenue'], report['sales_summary']['inventory']['count']])
                        writer.writerow(["Non-Inventory", report['sales_summary']['non_inventory']['revenue'], report['sales_summary']['non_inventory']['count']])
                        writer.writerow(["Rooms", report['sales_summary']['room']['revenue'], report['sales_summary']['room']['count']])
                        writer.writerow([])
                    
                    writer.writerow(["DAILY BREAKDOWN"])
                    writer.writerow(["Date", "Revenue", "Cost", "Profit", "Expenses", "Net Profit"])
                    for date, data in sorted(report['weekly_data'].items()):
                        writer.writerow([
                            date,
                            data.get('revenue', 0),
                            data.get('cost', 0),
                            data.get('profit', 0),
                            data.get('expenses', 0),
                            data.get('profit', 0) - data.get('expenses', 0)
                        ])
                    if report.get('expenses_detail'):
                        writer.writerow([])
                        writer.writerow(["WEEKLY EXPENSES DETAIL"])
                        writer.writerow(["Name", "Amount", "Date", "Notes"])
                        for exp in report['expenses_detail']:
                            writer.writerow([exp[1], exp[2], exp[3], exp[4] or ""])
                
                else:  # Daily Report
                    writer.writerow(["DAILY REPORT"])
                    writer.writerow(["Metric", "Value"])
                    writer.writerow(["Date", report.get('date', 'Unknown')])
                    writer.writerow(["Opening Balance", report.get('opening_balance', 0)])
                    writer.writerow(["Closing Balance", report.get('closing_balance', 0)])
                    writer.writerow(["Cash Sales", report.get('cash_sales', 0)])
                    writer.writerow(["Merchant Sales", report.get('merchant_sales', 0)])
                    writer.writerow(["Total Revenue", report.get('sales_revenue', 0)])
                    writer.writerow(["Gross Profit", report.get('sales_profit', 0)])
                    writer.writerow(["Total Expenses", report.get('expenses', 0)])
                    writer.writerow(["Net Profit", report.get('net_profit', 0)])
                    writer.writerow(["Occupied Rooms", f"{report.get('occupied', 0)}/{report.get('total_rooms', 0)}"])
                    writer.writerow(["Occupancy Rate", f"{report.get('occupied', 0)/report.get('total_rooms', 1)*100:.1f}%"])
                    
                    if report.get('is_saturday', False):
                        writer.writerow([])
                        writer.writerow(["SATURDAY COLLECTION"])
                        writer.writerow(["Cash Collected", report.get('cash_collected', 0)])
                        writer.writerow(["Merchant Collected", report.get('merchant_collected', 0)])
                        writer.writerow(["Total Collected", report.get('cash_collected', 0) + report.get('merchant_collected', 0)])
                    
                    if hasattr(self, 'current_sales_summary'):
                        writer.writerow([])
                        writer.writerow(["SALES BY TYPE"])
                        writer.writerow(["Type", "Revenue", "Count"])
                        writer.writerow(["Inventory", self.current_sales_summary['inventory']['revenue'], self.current_sales_summary['inventory']['count']])
                        writer.writerow(["Non-Inventory", self.current_sales_summary['non_inventory']['revenue'], self.current_sales_summary['non_inventory']['count']])
                        writer.writerow(["Rooms", self.current_sales_summary['room']['revenue'], self.current_sales_summary['room']['count']])
                    
                    if hasattr(self, 'current_expenses') and self.current_expenses:
                        writer.writerow([])
                        writer.writerow(["DAILY EXPENSES DETAIL"])
                        writer.writerow(["Name", "Amount", "Time", "Notes"])
                        for exp in self.current_expenses:
                            writer.writerow([exp[1], exp[2], exp[3], exp[4] or ""])
            
            self.show_message(f"✅ CSV exported: {filename}", "green")
            
        except Exception as e:
            self.show_message(f"Error exporting CSV: {str(e)}", "red")
    
    def export_word(self):
        """Export Word document"""
        if not hasattr(self, 'current_report'):
            self.show_message("Please generate a report first", "red")
            return
        
        if not HAS_DOCX:
            self.show_message("python-docx not installed. Run: pip install python-docx", "red")
            return
        
        try:
            report = self.current_report
            filename = f"Valley_View_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            
            doc = Document()
            
            # Title
            title = doc.add_heading('VALLEY VIEW MOTEL MANAGEMENT SYSTEM', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_heading(f'Report: {report.get("date", "Unknown")}', 1)
            doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            doc.add_paragraph()
            
            # Add content based on report type
            if report.get('date') == "All Data":
                doc.add_heading('LIFETIME TOTALS', 2)
                table = doc.add_table(rows=6, cols=2)
                table.style = 'Table Grid'
                cells = [
                    ("Total Sales", str(report.get('total_sales', 0))),
                    ("Total Revenue", f"UGX {report.get('total_revenue', 0):,.0f}"),
                    ("Total Profit", f"UGX {report.get('total_profit', 0):,.0f}"),
                    ("Total Expenses", f"UGX {report.get('total_expenses', 0):,.0f}"),
                    ("Net Profit", f"UGX {report.get('net_profit', 0):,.0f}")
                ]
                for i, (label, value) in enumerate(cells):
                    table.cell(i, 0).text = label
                    table.cell(i, 1).text = value
                
                if report.get('sales_summary'):
                    doc.add_heading('SALES BY TYPE', 2)
                    table = doc.add_table(rows=4, cols=3)
                    table.style = 'Table Grid'
                    table.cell(0, 0).text = "Type"
                    table.cell(0, 1).text = "Revenue"
                    table.cell(0, 2).text = "Count"
                    table.cell(1, 0).text = "Inventory"
                    table.cell(1, 1).text = f"UGX {report['sales_summary']['inventory']['revenue']:,.0f}"
                    table.cell(1, 2).text = str(report['sales_summary']['inventory']['count'])
                    table.cell(2, 0).text = "Non-Inventory"
                    table.cell(2, 1).text = f"UGX {report['sales_summary']['non_inventory']['revenue']:,.0f}"
                    table.cell(2, 2).text = str(report['sales_summary']['non_inventory']['count'])
                    table.cell(3, 0).text = "Rooms"
                    table.cell(3, 1).text = f"UGX {report['sales_summary']['room']['revenue']:,.0f}"
                    table.cell(3, 2).text = str(report['sales_summary']['room']['count'])
                
                # Expenses detail
                if report.get('expenses_detail'):
                    doc.add_heading('RECENT EXPENSES', 2)
                    table = doc.add_table(rows=len(report['expenses_detail']) + 1, cols=4)
                    table.style = 'Table Grid'
                    headers = ['Name', 'Amount', 'Date', 'Notes']
                    for i, h in enumerate(headers):
                        table.cell(0, i).text = h
                    for idx, exp in enumerate(report['expenses_detail'][:20], 1):
                        table.cell(idx, 0).text = exp[1]
                        table.cell(idx, 1).text = f"UGX {exp[2]:,.0f}"
                        table.cell(idx, 2).text = exp[3][:10] if exp[3] else ""
                        table.cell(idx, 3).text = exp[4] or ""
            
            elif report.get('date') == "Weekly":
                doc.add_heading('WEEKLY PERFORMANCE', 2)
                table = doc.add_table(rows=6, cols=2)
                table.style = 'Table Grid'
                cells = [
                    ("Total Revenue", f"UGX {report.get('revenue', 0):,.0f}"),
                    ("Cost of Goods", f"UGX {report.get('cost', 0):,.0f}"),
                    ("Gross Profit", f"UGX {report.get('gross_profit', 0):,.0f}"),
                    ("Total Expenses", f"UGX {report.get('expenses', 0):,.0f}"),
                    ("Net Profit", f"UGX {report.get('net_profit', 0):,.0f}")
                ]
                for i, (label, value) in enumerate(cells):
                    table.cell(i, 0).text = label
                    table.cell(i, 1).text = value
                
                if report.get('sales_summary'):
                    doc.add_heading('SALES BY TYPE', 2)
                    table = doc.add_table(rows=4, cols=3)
                    table.style = 'Table Grid'
                    table.cell(0, 0).text = "Type"
                    table.cell(0, 1).text = "Revenue"
                    table.cell(0, 2).text = "Count"
                    table.cell(1, 0).text = "Inventory"
                    table.cell(1, 1).text = f"UGX {report['sales_summary']['inventory']['revenue']:,.0f}"
                    table.cell(1, 2).text = str(report['sales_summary']['inventory']['count'])
                    table.cell(2, 0).text = "Non-Inventory"
                    table.cell(2, 1).text = f"UGX {report['sales_summary']['non_inventory']['revenue']:,.0f}"
                    table.cell(2, 2).text = str(report['sales_summary']['non_inventory']['count'])
                    table.cell(3, 0).text = "Rooms"
                    table.cell(3, 1).text = f"UGX {report['sales_summary']['room']['revenue']:,.0f}"
                    table.cell(3, 2).text = str(report['sales_summary']['room']['count'])
                
                doc.add_heading('DAILY BREAKDOWN', 2)
                if 'weekly_data' in report:
                    table = doc.add_table(rows=len(report['weekly_data']) + 1, cols=7)
                    table.style = 'Table Grid'
                    headers = ['Date', 'Revenue', 'Cost', 'Profit', 'Expenses', 'Net Profit']
                    for i, h in enumerate(headers):
                        table.cell(0, i).text = h
                    row_idx = 1
                    for date, data in sorted(report['weekly_data'].items()):
                        table.cell(row_idx, 0).text = date
                        table.cell(row_idx, 1).text = f"UGX {data.get('revenue', 0):,.0f}"
                        table.cell(row_idx, 2).text = f"UGX {data.get('cost', 0):,.0f}"
                        table.cell(row_idx, 3).text = f"UGX {data.get('profit', 0):,.0f}"
                        table.cell(row_idx, 4).text = f"UGX {data.get('expenses', 0):,.0f}"
                        table.cell(row_idx, 5).text = f"UGX {data.get('profit', 0) - data.get('expenses', 0):,.0f}"
                        row_idx += 1
                
                # Weekly expenses detail
                if report.get('expenses_detail'):
                    doc.add_heading('WEEKLY EXPENSES DETAIL', 2)
                    table = doc.add_table(rows=len(report['expenses_detail']) + 1, cols=4)
                    table.style = 'Table Grid'
                    headers = ['Name', 'Amount', 'Date', 'Notes']
                    for i, h in enumerate(headers):
                        table.cell(0, i).text = h
                    for idx, exp in enumerate(report['expenses_detail'][:20], 1):
                        table.cell(idx, 0).text = exp[1]
                        table.cell(idx, 1).text = f"UGX {exp[2]:,.0f}"
                        table.cell(idx, 2).text = exp[3][:10] if exp[3] else ""
                        table.cell(idx, 3).text = exp[4] or ""
            
            elif report.get('date') == "Balance Report":
                doc.add_heading('WEEKLY BALANCE REPORT', 2)
                if 'weekly_data' in report:
                    table = doc.add_table(rows=len(report['weekly_data']) + 1, cols=8)
                    table.style = 'Table Grid'
                    headers = ['Date', 'Opening', 'Revenue', 'Expenses', 'Closing', 'Cash', 'Merchant', 'Day']
                    for i, h in enumerate(headers):
                        table.cell(0, i).text = h
                    row_idx = 1
                    for data in report['weekly_data']:
                        table.cell(row_idx, 0).text = data["date"]
                        table.cell(row_idx, 1).text = f"UGX {data['opening_balance']:,.0f}"
                        table.cell(row_idx, 2).text = f"UGX {data['daily_revenue']:,.0f}"
                        table.cell(row_idx, 3).text = f"UGX {data['daily_expenses']:,.0f}"
                        table.cell(row_idx, 4).text = f"UGX {data['closing_balance']:,.0f}"
                        table.cell(row_idx, 5).text = f"UGX {data['cash_collected']:,.0f}"
                        table.cell(row_idx, 6).text = f"UGX {data['merchant_collected']:,.0f}"
                        table.cell(row_idx, 7).text = "SAT (Collection)" if data["is_saturday"] else ""
                        row_idx += 1
            
            else:  # Daily Report
                doc.add_heading('DAILY REPORT', 2)
                table = doc.add_table(rows=12, cols=2)
                table.style = 'Table Grid'
                cells = [
                    ("Date", report.get('date', 'Unknown')),
                    ("Opening Balance", f"UGX {report.get('opening_balance', 0):,.0f}"),
                    ("Closing Balance", f"UGX {report.get('closing_balance', 0):,.0f}"),
                    ("Cash Sales", f"UGX {report.get('cash_sales', 0):,.0f}"),
                    ("Merchant Sales", f"UGX {report.get('merchant_sales', 0):,.0f}"),
                    ("Total Revenue", f"UGX {report.get('sales_revenue', 0):,.0f}"),
                    ("Gross Profit", f"UGX {report.get('sales_profit', 0):,.0f}"),
                    ("Total Expenses", f"UGX {report.get('expenses', 0):,.0f}"),
                    ("Net Profit", f"UGX {report.get('net_profit', 0):,.0f}"),
                    ("Occupied Rooms", f"{report.get('occupied', 0)}/{report.get('total_rooms', 0)}"),
                    ("Occupancy Rate", f"{report.get('occupied', 0)/report.get('total_rooms', 1)*100:.1f}%")
                ]
                for i, (label, value) in enumerate(cells):
                    table.cell(i, 0).text = label
                    table.cell(i, 1).text = value
                
                if hasattr(self, 'current_sales_summary'):
                    doc.add_heading('SALES BY TYPE', 2)
                    table = doc.add_table(rows=4, cols=3)
                    table.style = 'Table Grid'
                    table.cell(0, 0).text = "Type"
                    table.cell(0, 1).text = "Revenue"
                    table.cell(0, 2).text = "Count"
                    table.cell(1, 0).text = "Inventory"
                    table.cell(1, 1).text = f"UGX {self.current_sales_summary['inventory']['revenue']:,.0f}"
                    table.cell(1, 2).text = str(self.current_sales_summary['inventory']['count'])
                    table.cell(2, 0).text = "Non-Inventory"
                    table.cell(2, 1).text = f"UGX {self.current_sales_summary['non_inventory']['revenue']:,.0f}"
                    table.cell(2, 2).text = str(self.current_sales_summary['non_inventory']['count'])
                    table.cell(3, 0).text = "Rooms"
                    table.cell(3, 1).text = f"UGX {self.current_sales_summary['room']['revenue']:,.0f}"
                    table.cell(3, 2).text = str(self.current_sales_summary['room']['count'])
                
                # Daily expenses detail
                if hasattr(self, 'current_expenses') and self.current_expenses:
                    doc.add_heading('DAILY EXPENSES DETAIL', 2)
                    table = doc.add_table(rows=len(self.current_expenses) + 1, cols=4)
                    table.style = 'Table Grid'
                    headers = ['Name', 'Amount', 'Time', 'Notes']
                    for i, h in enumerate(headers):
                        table.cell(0, i).text = h
                    for idx, exp in enumerate(self.current_expenses, 1):
                        table.cell(idx, 0).text = exp[1]
                        table.cell(idx, 1).text = f"UGX {exp[2]:,.0f}"
                        table.cell(idx, 2).text = exp[3][:16] if exp[3] else ""
                        table.cell(idx, 3).text = exp[4] or ""
                
                if report.get('is_saturday', False):
                    doc.add_heading('SATURDAY COLLECTION', 2)
                    table = doc.add_table(rows=4, cols=2)
                    table.style = 'Table Grid'
                    cells = [
                        ("Cash Collected", f"UGX {report.get('cash_collected', 0):,.0f}"),
                        ("Merchant Collected", f"UGX {report.get('merchant_collected', 0):,.0f}"),
                        ("Total Collected", f"UGX {report.get('cash_collected', 0) + report.get('merchant_collected', 0):,.0f}")
                    ]
                    for i, (label, value) in enumerate(cells):
                        table.cell(i, 0).text = label
                        table.cell(i, 1).text = value
            
            # Footer
            doc.add_paragraph()
            doc.add_paragraph('-' * 50)
            doc.add_paragraph('End of Report')
            doc.add_paragraph(f'Generated by Valley View Motel Management System v1.0')
            
            doc.save(filename)
            self.show_message(f"✅ Word document exported: {filename}", "green")
            
        except Exception as e:
            self.show_message(f"Error exporting Word: {str(e)}", "red")
    
    def export_pdf(self):
        """Export PDF report"""
        if not hasattr(self, 'current_report'):
            self.show_message("Please generate a report first", "red")
            return
        
        if not HAS_REPORTLAB:
            self.show_message("reportlab not installed. Run: pip install reportlab", "red")
            return
        
        try:
            report = self.current_report
            filename = f"Valley_View_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            
            doc = SimpleDocTemplate(filename, pagesize=A4)
            styles = getSampleStyleSheet()
            story = []
            
            # Title
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                textColor=colors.HexColor('#1f538d'),
                alignment=1,
                spaceAfter=30
            )
            story.append(Paragraph("VALLEY VIEW MOTEL", title_style))
            story.append(Spacer(1, 0.2*inch))
            
            # Subtitle
            sub_style = ParagraphStyle(
                'CustomSub',
                parent=styles['Normal'],
                fontSize=16,
                alignment=1
            )
            story.append(Paragraph(f"Report: {report.get('date', 'Unknown')}", sub_style))
            story.append(Paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", sub_style))
            story.append(Spacer(1, 0.3*inch))
            
            # Add content based on report type
            if report.get('date') == "All Data":
                story.append(Paragraph("LIFETIME TOTALS", styles['Heading2']))
                data = [
                    ["Metric", "Value"],
                    ["Total Sales", str(report.get('total_sales', 0))],
                    ["Total Revenue", f"UGX {report.get('total_revenue', 0):,.0f}"],
                    ["Total Profit", f"UGX {report.get('total_profit', 0):,.0f}"],
                    ["Total Expenses", f"UGX {report.get('total_expenses', 0):,.0f}"],
                    ["Net Profit", f"UGX {report.get('net_profit', 0):,.0f}"]
                ]
                table = Table(data, colWidths=[3*inch, 3*inch])
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 14),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                story.append(table)
                
                # Sales by type
                if report.get('sales_summary'):
                    story.append(Spacer(1, 0.3*inch))
                    story.append(Paragraph("SALES BY TYPE", styles['Heading2']))
                    data = [
                        ["Type", "Revenue", "Count"],
                        ["Inventory", f"UGX {report['sales_summary']['inventory']['revenue']:,.0f}", str(report['sales_summary']['inventory']['count'])],
                        ["Non-Inventory", f"UGX {report['sales_summary']['non_inventory']['revenue']:,.0f}", str(report['sales_summary']['non_inventory']['count'])],
                        ["Rooms", f"UGX {report['sales_summary']['room']['revenue']:,.0f}", str(report['sales_summary']['room']['count'])]
                    ]
                    table = Table(data, colWidths=[2*inch, 2*inch, 1*inch])
                    table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 10),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black)
                    ]))
                    story.append(table)
                
                # Expenses detail
                if report.get('expenses_detail'):
                    story.append(Spacer(1, 0.3*inch))
                    story.append(Paragraph("RECENT EXPENSES", styles['Heading2']))
                    data = [["Name", "Amount", "Date", "Notes"]]
                    for exp in report['expenses_detail'][:15]:
                        data.append([
                            exp[1],
                            f"UGX {exp[2]:,.0f}",
                            exp[3][:10] if exp[3] else "",
                            exp[4] or ""
                        ])
                    table = Table(data, colWidths=[2*inch, 1.5*inch, 1.5*inch, 1.5*inch])
                    table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 10),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black),
                        ('FONTSIZE', (0, 1), (-1, -1), 8)
                    ]))
                    story.append(table)
            
            elif report.get('date') == "Weekly":
                story.append(Paragraph("WEEKLY PERFORMANCE", styles['Heading2']))
                data = [
                    ["Metric", "Value"],
                    ["Total Revenue", f"UGX {report.get('revenue', 0):,.0f}"],
                    ["Cost of Goods", f"UGX {report.get('cost', 0):,.0f}"],
                    ["Gross Profit", f"UGX {report.get('gross_profit', 0):,.0f}"],
                    ["Total Expenses", f"UGX {report.get('expenses', 0):,.0f}"],
                    ["Net Profit", f"UGX {report.get('net_profit', 0):,.0f}"]
                ]
                table = Table(data, colWidths=[3*inch, 3*inch])
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 14),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                story.append(table)
                story.append(Spacer(1, 0.3*inch))
                
                # Sales by type
                if report.get('sales_summary'):
                    story.append(Paragraph("SALES BY TYPE", styles['Heading2']))
                    data = [
                        ["Type", "Revenue", "Count"],
                        ["Inventory", f"UGX {report['sales_summary']['inventory']['revenue']:,.0f}", str(report['sales_summary']['inventory']['count'])],
                        ["Non-Inventory", f"UGX {report['sales_summary']['non_inventory']['revenue']:,.0f}", str(report['sales_summary']['non_inventory']['count'])],
                        ["Rooms", f"UGX {report['sales_summary']['room']['revenue']:,.0f}", str(report['sales_summary']['room']['count'])]
                    ]
                    table = Table(data, colWidths=[2*inch, 2*inch, 1*inch])
                    table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 10),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black)
                    ]))
                    story.append(table)
                
                # Daily breakdown
                story.append(Spacer(1, 0.3*inch))
                story.append(Paragraph("DAILY BREAKDOWN", styles['Heading2']))
                if 'weekly_data' in report:
                    data = [["Date", "Revenue", "Cost", "Profit", "Expenses", "Net Profit"]]
                    for date, d in sorted(report['weekly_data'].items()):
                        data.append([
                            date,
                            f"UGX {d.get('revenue', 0):,.0f}",
                            f"UGX {d.get('cost', 0):,.0f}",
                            f"UGX {d.get('profit', 0):,.0f}",
                            f"UGX {d.get('expenses', 0):,.0f}",
                            f"UGX {d.get('profit', 0) - d.get('expenses', 0):,.0f}"
                        ])
                    table = Table(data, colWidths=[1.2*inch, 1.2*inch, 1.2*inch, 1.2*inch, 1.2*inch, 1.2*inch])
                    table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 10),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black),
                        ('FONTSIZE', (0, 1), (-1, -1), 9)
                    ]))
                    story.append(table)
                
                # Weekly expenses detail
                if report.get('expenses_detail'):
                    story.append(Spacer(1, 0.3*inch))
                    story.append(Paragraph("WEEKLY EXPENSES DETAIL", styles['Heading2']))
                    data = [["Name", "Amount", "Date", "Notes"]]
                    for exp in report['expenses_detail'][:15]:
                        data.append([
                            exp[1],
                            f"UGX {exp[2]:,.0f}",
                            exp[3][:10] if exp[3] else "",
                            exp[4] or ""
                        ])
                    table = Table(data, colWidths=[2*inch, 1.5*inch, 1.5*inch, 1.5*inch])
                    table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 10),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black),
                        ('FONTSIZE', (0, 1), (-1, -1), 8)
                    ]))
                    story.append(table)
            
            elif report.get('date') == "Balance Report":
                story.append(Paragraph("WEEKLY BALANCE REPORT", styles['Heading2']))
                if 'weekly_data' in report:
                    data = [["Date", "Opening", "Revenue", "Expenses", "Closing", "Cash", "Merchant", "Day"]]
                    for d in report['weekly_data']:
                        data.append([
                            d["date"],
                            f"UGX {d['opening_balance']:,.0f}",
                            f"UGX {d['daily_revenue']:,.0f}",
                            f"UGX {d['daily_expenses']:,.0f}",
                            f"UGX {d['closing_balance']:,.0f}",
                            f"UGX {d['cash_collected']:,.0f}",
                            f"UGX {d['merchant_collected']:,.0f}",
                            "SAT (Collection)" if d["is_saturday"] else ""
                        ])
                    table = Table(data, colWidths=[1*inch, 1*inch, 1*inch, 1*inch, 1*inch, 1*inch, 1*inch, 1*inch])
                    table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 9),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black),
                        ('FONTSIZE', (0, 1), (-1, -1), 8)
                    ]))
                    story.append(table)
            
            else:  # Daily Report
                story.append(Paragraph("DAILY REPORT", styles['Heading2']))
                data = [
                    ["Metric", "Value"],
                    ["Date", report.get('date', 'Unknown')],
                    ["Opening Balance", f"UGX {report.get('opening_balance', 0):,.0f}"],
                    ["Closing Balance", f"UGX {report.get('closing_balance', 0):,.0f}"],
                    ["Cash Sales", f"UGX {report.get('cash_sales', 0):,.0f}"],
                    ["Merchant Sales", f"UGX {report.get('merchant_sales', 0):,.0f}"],
                    ["Total Revenue", f"UGX {report.get('sales_revenue', 0):,.0f}"],
                    ["Gross Profit", f"UGX {report.get('sales_profit', 0):,.0f}"],
                    ["Total Expenses", f"UGX {report.get('expenses', 0):,.0f}"],
                    ["Net Profit", f"UGX {report.get('net_profit', 0):,.0f}"],
                    ["Occupied Rooms", f"{report.get('occupied', 0)}/{report.get('total_rooms', 0)}"],
                    ["Occupancy Rate", f"{report.get('occupied', 0)/report.get('total_rooms', 1)*100:.1f}%"]
                ]
                table = Table(data, colWidths=[3*inch, 3*inch])
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 14),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                story.append(table)
                
                # Sales by type
                if hasattr(self, 'current_sales_summary'):
                    story.append(Spacer(1, 0.3*inch))
                    story.append(Paragraph("SALES BY TYPE", styles['Heading2']))
                    data = [
                        ["Type", "Revenue", "Count"],
                        ["Inventory", f"UGX {self.current_sales_summary['inventory']['revenue']:,.0f}", str(self.current_sales_summary['inventory']['count'])],
                        ["Non-Inventory", f"UGX {self.current_sales_summary['non_inventory']['revenue']:,.0f}", str(self.current_sales_summary['non_inventory']['count'])],
                        ["Rooms", f"UGX {self.current_sales_summary['room']['revenue']:,.0f}", str(self.current_sales_summary['room']['count'])]
                    ]
                    table = Table(data, colWidths=[2*inch, 2*inch, 1*inch])
                    table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 10),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black)
                    ]))
                    story.append(table)
                
                # Daily expenses detail
                if hasattr(self, 'current_expenses') and self.current_expenses:
                    story.append(Spacer(1, 0.3*inch))
                    story.append(Paragraph("DAILY EXPENSES DETAIL", styles['Heading2']))
                    data = [["Name", "Amount", "Time", "Notes"]]
                    for exp in self.current_expenses[:15]:
                        data.append([
                            exp[1],
                            f"UGX {exp[2]:,.0f}",
                            exp[3][:16] if exp[3] else "",
                            exp[4] or ""
                        ])
                    table = Table(data, colWidths=[2*inch, 1.5*inch, 1.5*inch, 1.5*inch])
                    table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 10),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black),
                        ('FONTSIZE', (0, 1), (-1, -1), 8)
                    ]))
                    story.append(table)
                
                if report.get('is_saturday', False):
                    story.append(Spacer(1, 0.3*inch))
                    story.append(Paragraph("SATURDAY COLLECTION", styles['Heading2']))
                    data = [
                        ["Cash Collected", f"UGX {report.get('cash_collected', 0):,.0f}"],
                        ["Merchant Collected", f"UGX {report.get('merchant_collected', 0):,.0f}"],
                        ["Total Collected", f"UGX {report.get('cash_collected', 0) + report.get('merchant_collected', 0):,.0f}"]
                    ]
                    table = Table(data, colWidths=[3*inch, 3*inch])
                    table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black)
                    ]))
                    story.append(table)
            
            # Footer
            story.append(Spacer(1, 0.5*inch))
            story.append(Paragraph("End of Report", styles['Normal']))
            story.append(Paragraph("Generated by Valley View Motel Management System v1.0", styles['Normal']))
            
            doc.build(story)
            self.show_message(f"✅ PDF exported: {filename}", "green")
            
        except Exception as e:
            self.show_message(f"Error exporting PDF: {str(e)}", "red")
    
    def show_message(self, msg, color="green"):
        """Show a message"""
        self.clear_report()
        font_size = get_font_size()
        ctk.CTkLabel(
            self.report_frame,
            text=msg,
            text_color=color,
            font=get_font(font_size + 2)
        ).pack(pady=20)